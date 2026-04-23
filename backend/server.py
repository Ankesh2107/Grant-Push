"""GrantPulse backend.

Autonomous agent platform for grant discovery and proposal drafting.
"""
from __future__ import annotations

import asyncio
import io
import json
import logging
import os
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

import bcrypt
import httpx
import jwt
import requests
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from dotenv import load_dotenv
from fastapi import (APIRouter, Depends, FastAPI, File, Header, HTTPException,
                     Query, Request, UploadFile)
from fastapi.responses import Response
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, EmailStr, Field
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from starlette.middleware.cors import CORSMiddleware



ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("grantpulse")

# ---- Config ----
MONGO_URL = os.environ["MONGO_URL"]
DB_NAME = os.environ["DB_NAME"]
JWT_SECRET = os.environ.get("JWT_SECRET", "dev-secret")
JWT_ALG = "HS256"
EMERGENT_KEY = os.environ["EMERGENT_LLM_KEY"]
STRIPE_API_KEY = os.environ["STRIPE_API_KEY"]
APP_NAME = "grantpulse"
STORAGE_URL = "https://integrations.emergentagent.com/objstore/api/v1/storage"

# Pricing (backend-defined, never trust frontend)
PRICE_PACKAGES = {"pro_monthly": 199.00}

# ---- DB ----
client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]

app = FastAPI(title="GrantPulse")
api = APIRouter(prefix="/api")

# ---- Storage helpers ----
_storage_key: Optional[str] = None


def init_storage() -> str:
    global _storage_key
    if _storage_key:
        return _storage_key
    resp = requests.post(
        f"{STORAGE_URL}/init",
        json={"emergent_key": EMERGENT_KEY},
        timeout=30,
    )
    resp.raise_for_status()
    _storage_key = resp.json()["storage_key"]
    return _storage_key


def put_object(path: str, data: bytes, content_type: str) -> dict:
    key = init_storage()
    resp = requests.put(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key, "Content-Type": content_type},
        data=data,
        timeout=120,
    )
    resp.raise_for_status()
    return resp.json()


def get_object(path: str) -> tuple[bytes, str]:
    key = init_storage()
    resp = requests.get(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key},
        timeout=60,
    )
    resp.raise_for_status()
    return resp.content, resp.headers.get("Content-Type", "application/octet-stream")


# ---- Auth helpers ----
def hash_password(pw: str) -> str:
    return bcrypt.hashpw(pw.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def verify_password(pw: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(pw.encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False


def make_token(user_id: str) -> str:
    payload = {
        "sub": user_id,
        "iat": datetime.now(timezone.utc),
        "exp": datetime.now(timezone.utc) + timedelta(days=30),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)


async def current_user(authorization: Optional[str] = Header(None)) -> dict:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(401, "Missing token")
    token = authorization.split(" ", 1)[1]
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
    except jwt.InvalidTokenError:
        raise HTTPException(401, "Invalid token")
    user = await db.users.find_one({"id": payload["sub"]}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(401, "User not found")
    return user


# ---- Models ----
class SignupReq(BaseModel):
    email: EmailStr
    password: str = Field(min_length=6)
    company_name: Optional[str] = None


class LoginReq(BaseModel):
    email: EmailStr
    password: str


class PersonaReq(BaseModel):
    company_name: str
    capabilities: List[str] = []
    past_performance: List[Dict[str, Any]] = []
    technical_keywords: List[str] = []
    geographic_focus: List[str] = []
    narrative: Optional[str] = ""


class DraftUpdateReq(BaseModel):
    content: str


class CheckoutReq(BaseModel):
    package_id: str
    origin_url: str


# ---- Utilities ----
def utcnow_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def new_id() -> str:
    return str(uuid.uuid4())


def clean_mongo(doc: dict) -> dict:
    doc.pop("_id", None)
    return doc


# ---- LLM helpers ----
from openai import AsyncOpenAI
import stripe
import os

stripe.api_key = STRIPE_API_KEY

# ---- LLM helpers ----
from openai import AsyncOpenAI
import stripe
import os
import json

stripe.api_key = STRIPE_API_KEY

async def score_grant_vs_persona(grant: dict, persona: dict) -> dict:
    system = "You are GrantPulse's vetting engine. You analyze how well a company persona matches an RFP/grant opportunity. Respond with STRICT JSON only."
    prompt = f"Analyze this match. Output JSON with keys: score (0-100 integer), reasoning (2-3 sentences), summary (3-sentence executive summary).\n\nCOMPANY PERSONA:\n{json.dumps(persona, indent=2)}\n\nGRANT/RFP:\nTitle: {grant.get('title')}\nAgency: {grant.get('agency')}\nDescription: {grant.get('description','')[:3000]}"
    try:
        # Pointing the OpenAI library to Groq's free servers
        client = AsyncOpenAI(api_key=os.environ.get("GROQ_API_KEY"), base_url="https://api.groq.com/openai/v1")
        res = await client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}], temperature=0.2)
        data = json.loads(res.choices[0].message.content.strip("`").replace("json", "", 1).strip())
        return {"score": int(data.get("score", 0)), "reasoning": str(data.get("reasoning", ""))[:1000], "summary": str(data.get("summary", ""))[:600]}
    except Exception as e:
        return {"score": 0, "reasoning": "Scoring failed", "summary": "N/A"}

async def draft_proposal(grant: dict, persona: dict, past_proposals_text: list) -> str:
    system = "You are an elite proposal drafter. Mirror the tone of past proposals."
    prompt = f"Draft a full RFP response.\nCOMPANY PERSONA: {json.dumps(persona)}\nOPPORTUNITY: {grant.get('title')}\nDesc: {grant.get('description','')[:4000]}"
    client = AsyncOpenAI(api_key=os.environ.get("GROQ_API_KEY"), base_url="https://api.groq.com/openai/v1")
    res = await client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}], temperature=0.5)
    return res.choices[0].message.content

# ---- Grant fetching (Grants.gov public API) ----
GRANTS_API = "https://api.grants.gov/v1/api/search2"


async def fetch_grants_gov(keywords: str = "", rows: int = 25) -> List[dict]:
    """Query the official Grants.gov search API. Returns normalized grant dicts."""
    body = {"rows": rows, "keyword": keywords, "oppStatuses": "forecasted|posted"}
    try:
        async with httpx.AsyncClient(timeout=30) as hc:
            r = await hc.post(GRANTS_API, json=body)
            r.raise_for_status()
            data = r.json()
    except Exception as e:
        logger.error(f"Grants.gov fetch failed: {e}")
        return []
    out = []
    hits = (data.get("data", {}) or {}).get("oppHits", []) or []
    for h in hits:
        out.append({
            "source": "grants.gov",
            "external_id": str(h.get("id") or h.get("number") or ""),
            "title": h.get("title", ""),
            "agency": h.get("agency", "") or h.get("agencyName", ""),
            "category": h.get("oppStatus", ""),
            "funding_amount": h.get("awardCeiling") or h.get("awardFloor") or "",
            "deadline": h.get("closeDate", ""),
            "posted_date": h.get("openDate", ""),
            "description": h.get("description") or h.get("title", ""),
            "url": f"https://www.grants.gov/search-results-detail/{h.get('id','')}" if h.get("id") else "",
        })
    return out


# ---- Routes: Auth ----
@api.get("/")
async def health():
    return {"status": "ok", "service": "grantpulse"}


@api.post("/auth/signup")
async def signup(req: SignupReq):
    existing = await db.users.find_one({"email": req.email})
    if existing:
        raise HTTPException(400, "Email already registered")
    uid = new_id()
    doc = {
        "id": uid,
        "email": req.email,
        "password": hash_password(req.password),
        "role": "admin" if req.email == "admin@grantpulse.io" else "client",
        "plan": "free",
        "free_leads_used": 0,
        "free_leads_period_start": utcnow_iso(),
        "created_at": utcnow_iso(),
        "company_name": req.company_name or "",
    }
    await db.users.insert_one(doc)
    token = make_token(uid)
    user = {k: v for k, v in doc.items() if k not in ("password", "_id")}
    return {"token": token, "user": user}


@api.post("/auth/login")
async def login(req: LoginReq):
    user = await db.users.find_one({"email": req.email})
    if not user or not verify_password(req.password, user.get("password", "")):
        raise HTTPException(401, "Invalid credentials")
    token = make_token(user["id"])
    user = clean_mongo(user)
    user.pop("password", None)
    return {"token": token, "user": user}


@api.get("/auth/me")
async def me(user: dict = Depends(current_user)):
    return user


# ---- Routes: Persona ----
@api.get("/persona")
async def get_persona(user: dict = Depends(current_user)):
    doc = await db.personas.find_one({"user_id": user["id"]}, {"_id": 0})
    return doc or {}


@api.put("/persona")
async def save_persona(req: PersonaReq, user: dict = Depends(current_user)):
    doc = req.model_dump()
    doc["user_id"] = user["id"]
    doc["updated_at"] = utcnow_iso()
    await db.personas.update_one(
        {"user_id": user["id"]}, {"$set": doc}, upsert=True
    )
    return {"ok": True}


# ---- Routes: Grants / Scout ----
async def _run_scout_for_user(user: dict) -> dict:
    persona = await db.personas.find_one({"user_id": user["id"]}, {"_id": 0})
    if not persona:
        return {"scanned": 0, "created": 0, "error": "persona_not_set"}
    keywords = " ".join(persona.get("technical_keywords", [])[:5])
    grants = await fetch_grants_gov(keywords=keywords, rows=15)
    created = 0
    matched = 0
    for g in grants:
        # dedupe by user + external_id
        key = {"user_id": user["id"], "external_id": g["external_id"]}
        if await db.grants.find_one(key):
            continue
        gid = new_id()
        # Score
        try:
            scored = await score_grant_vs_persona(g, persona)
        except Exception as e:
            logger.error(f"score error: {e}")
            scored = {"score": 0, "reasoning": "scoring failed", "summary": g["description"][:300]}
        doc = {
            "id": gid,
            "user_id": user["id"],
            **g,
            "pow_score": scored["score"],
            "reasoning": scored["reasoning"],
            "summary": scored["summary"],
            "stage": "matched" if scored["score"] >= 85 else "matched",  # noqa
            "high_match": scored["score"] >= 85,
            "draft_id": None,
            "created_at": utcnow_iso(),
        }
        await db.grants.insert_one(doc)
        created += 1
        if scored["score"] >= 85:
            matched += 1
            # In-app notification
            await db.notifications.insert_one({
                "id": new_id(),
                "user_id": user["id"],
                "type": "high_match",
                "grant_id": gid,
                "message": f"NEW HIGH-MATCH: {g['title'][:80]} — PoW {scored['score']}%",
                "created_at": utcnow_iso(),
                "read": False,
            })
    return {"scanned": len(grants), "created": created, "high_match": matched}


@api.post("/scout/run")
async def scout_run(user: dict = Depends(current_user)):
    # Free tier gate: 1 lead / month
    if user.get("plan", "free") == "free":
        used = user.get("free_leads_used", 0)
        if used >= 1:
            # Still allow the scan but only surface 1 lead; we enforce on draft side too
            pass
    result = await _run_scout_for_user(user)
    return result


@api.get("/grants")
async def list_grants(
    stage: Optional[str] = None,
    user: dict = Depends(current_user),
):
    q: dict = {"user_id": user["id"]}
    if stage:
        q["stage"] = stage
    docs = await db.grants.find(q, {"_id": 0}).sort("pow_score", -1).to_list(500)
    # Free-tier gating: only expose 1 grant
    if user.get("plan", "free") == "free":
        # expose only top-1 high-match, rest are locked previews
        for i, d in enumerate(docs):
            if i >= 1:
                d["locked"] = True
    return docs


@api.get("/grants/{gid}")
async def get_grant(gid: str, user: dict = Depends(current_user)):
    doc = await db.grants.find_one({"id": gid, "user_id": user["id"]}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Not found")
    return doc


@api.post("/grants/{gid}/stage")
async def move_stage(gid: str, stage: str, user: dict = Depends(current_user)):
    if stage not in {"matched", "drafting", "review", "submitted"}:
        raise HTTPException(400, "Bad stage")
    res = await db.grants.update_one(
        {"id": gid, "user_id": user["id"]},
        {"$set": {"stage": stage, "updated_at": utcnow_iso()}},
    )
    if res.matched_count == 0:
        raise HTTPException(404, "Not found")
    return {"ok": True}


# ---- Routes: Drafts ----
@api.post("/grants/{gid}/draft")
async def create_draft(gid: str, user: dict = Depends(current_user)):
    # Gate drafting to paid plan
    if user.get("plan", "free") != "pro":
        raise HTTPException(402, "AI drafting requires Pro plan")
    grant = await db.grants.find_one({"id": gid, "user_id": user["id"]}, {"_id": 0})
    if not grant:
        raise HTTPException(404, "Grant not found")
    persona = await db.personas.find_one({"user_id": user["id"]}, {"_id": 0}) or {}
    past = await db.vault_files.find(
        {"user_id": user["id"], "is_deleted": False}, {"_id": 0}
    ).to_list(10)
    past_texts = [p.get("extracted_text", "")[:6000] for p in past if p.get("extracted_text")]
    did = new_id()
    doc = {
        "id": did,
        "user_id": user["id"],
        "grant_id": gid,
        "content": "",
        "status": "generating",
        "error": None,
        "created_at": utcnow_iso(),
        "approved_at": None,
        "submitted_at": None,
    }
    await db.drafts.insert_one(doc)
    await db.grants.update_one(
        {"id": gid},
        {"$set": {"stage": "drafting", "draft_id": did, "updated_at": utcnow_iso()}},
    )

    async def _generate():
        try:
            content = await draft_proposal(grant, persona, past_texts)
            await db.drafts.update_one(
                {"id": did},
                {"$set": {"content": content, "status": "review", "updated_at": utcnow_iso()}},
            )
            await db.grants.update_one(
                {"id": gid}, {"$set": {"stage": "review", "updated_at": utcnow_iso()}}
            )
        except Exception as exc:
            msg = str(exc)
            logger.error(f"draft gen failed {did}: {msg}")
            friendly = msg[:500]
            if "budget" in msg.lower():
                friendly = (
                    "LLM budget exhausted on the Emergent universal key. "
                    "Top up your balance (Profile → Universal Key → Add Balance) and retry."
                )
            await db.drafts.update_one(
                {"id": did},
                {"$set": {"status": "failed", "error": friendly, "updated_at": utcnow_iso()}},
            )

    asyncio.create_task(_generate())
    return clean_mongo({k: v for k, v in doc.items()})


@api.get("/drafts/{did}")
async def get_draft(did: str, user: dict = Depends(current_user)):
    d = await db.drafts.find_one({"id": did, "user_id": user["id"]}, {"_id": 0})
    if not d:
        raise HTTPException(404, "Not found")
    return d


@api.put("/drafts/{did}")
async def update_draft(did: str, req: DraftUpdateReq, user: dict = Depends(current_user)):
    res = await db.drafts.update_one(
        {"id": did, "user_id": user["id"]},
        {"$set": {"content": req.content, "updated_at": utcnow_iso()}},
    )
    if res.matched_count == 0:
        raise HTTPException(404, "Not found")
    return {"ok": True}


@api.post("/drafts/{did}/approve")
async def approve_draft(did: str, user: dict = Depends(current_user)):
    """Trust Boundary: explicit user action required to mark as submitted."""
    d = await db.drafts.find_one({"id": did, "user_id": user["id"]})
    if not d:
        raise HTTPException(404, "Not found")
    now = utcnow_iso()
    await db.drafts.update_one(
        {"id": did},
        {"$set": {"status": "submitted", "approved_at": now, "submitted_at": now}},
    )
    await db.grants.update_one(
        {"id": d["grant_id"]},
        {"$set": {"stage": "submitted", "updated_at": now}},
    )
    return {"ok": True, "submitted_at": now}


def _build_docx(draft: dict, grant: dict, persona: dict) -> bytes:
    """Render a draft into a formatted DOCX document."""
    doc = Document()
    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover block
    title = doc.add_heading(grant.get("title") or "RFP Response", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
    meta = doc.add_paragraph()
    meta.add_run(f"Submitted by: {persona.get('company_name') or '—'}\n").bold = True
    meta.add_run(f"Agency: {grant.get('agency') or '—'}\n")
    meta.add_run(f"Opportunity ID: {grant.get('external_id') or '—'}\n")
    meta.add_run(f"Deadline: {grant.get('deadline') or '—'}\n")
    meta.add_run(f"Probability of Win: {grant.get('pow_score', 0)}%\n")
    doc.add_paragraph("")

    # Body: parse markdown-ish content (# H1, ## H2, blank lines = paragraph breaks)
    body = draft.get("content") or ""
    for line in body.splitlines():
        stripped = line.rstrip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    # Footer note
    foot = doc.add_paragraph()
    foot.add_run(
        f"\n\nDrafted by GrantPulse · Status: {draft.get('status','').upper()} · Draft ID: {draft.get('id')}"
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@api.get("/drafts/{did}/export")
async def export_draft_docx(did: str, user: dict = Depends(current_user)):
    draft = await db.drafts.find_one({"id": did, "user_id": user["id"]}, {"_id": 0})
    if not draft:
        raise HTTPException(404, "Not found")
    if draft.get("status") == "generating":
        raise HTTPException(409, "Draft is still generating")
    if not (draft.get("content") or "").strip():
        raise HTTPException(400, "Draft has no content to export")
    grant = await db.grants.find_one({"id": draft["grant_id"]}, {"_id": 0}) or {}
    persona = await db.personas.find_one({"user_id": user["id"]}, {"_id": 0}) or {}
    data = _build_docx(draft, grant, persona)
    safe_title = "".join(c for c in (grant.get("title") or "grantpulse-draft") if c.isalnum() or c in " -_")[:60].strip() or "grantpulse-draft"
    filename = f"{safe_title}.docx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---- Routes: Vault (PDF upload) ----
@api.post("/vault/upload")
async def upload_vault(
    file: UploadFile = File(...), user: dict = Depends(current_user)
):
    if not (file.filename or "").lower().endswith(".pdf"):
        raise HTTPException(400, "PDF only")
    data = await file.read()
    # Extract text
    text = ""
    try:
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception as e:
        logger.warning(f"PDF extract failed: {e}")
    path = f"{APP_NAME}/vault/{user['id']}/{new_id()}.pdf"
    try:
        res = put_object(path, data, "application/pdf")
        stored_path = res["path"]
        size = res.get("size", len(data))
    except Exception as e:
        logger.error(f"storage failed: {e}")
        stored_path = path
        size = len(data)
    fid = new_id()
    doc = {
        "id": fid,
        "user_id": user["id"],
        "storage_path": stored_path,
        "original_filename": file.filename,
        "content_type": "application/pdf",
        "size": size,
        "extracted_text": text[:50000],
        "is_deleted": False,
        "created_at": utcnow_iso(),
    }
    await db.vault_files.insert_one(doc)
    return clean_mongo({k: v for k, v in doc.items() if k != "extracted_text"})


@api.get("/vault")
async def list_vault(user: dict = Depends(current_user)):
    docs = await db.vault_files.find(
        {"user_id": user["id"], "is_deleted": False},
        {"_id": 0, "extracted_text": 0},
    ).sort("created_at", -1).to_list(100)
    return docs


@api.delete("/vault/{fid}")
async def delete_vault(fid: str, user: dict = Depends(current_user)):
    await db.vault_files.update_one(
        {"id": fid, "user_id": user["id"]}, {"$set": {"is_deleted": True}}
    )
    return {"ok": True}


# ---- Routes: Notifications ----
@api.get("/notifications")
async def list_notifications(user: dict = Depends(current_user)):
    docs = await db.notifications.find(
        {"user_id": user["id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    return docs


@api.post("/notifications/read-all")
async def read_all(user: dict = Depends(current_user)):
    await db.notifications.update_many(
        {"user_id": user["id"]}, {"$set": {"read": True}}
    )
    return {"ok": True}


# ---- Routes: Dashboard metrics ----
@api.get("/dashboard/metrics")
async def dashboard_metrics(user: dict = Depends(current_user)):
    uid = user["id"]
    total = await db.grants.count_documents({"user_id": uid})
    matched = await db.grants.count_documents({"user_id": uid, "stage": "matched"})
    drafting = await db.grants.count_documents({"user_id": uid, "stage": "drafting"})
    review = await db.grants.count_documents({"user_id": uid, "stage": "review"})
    submitted = await db.grants.count_documents({"user_id": uid, "stage": "submitted"})
    high = await db.grants.count_documents({"user_id": uid, "high_match": True})
    unread = await db.notifications.count_documents({"user_id": uid, "read": False})
    return {
        "total_grants": total,
        "matched": matched,
        "drafting": drafting,
        "review": review,
        "submitted": submitted,
        "high_match": high,
        "unread_notifications": unread,
        "active_scouts": 1 if await db.personas.find_one({"user_id": uid}) else 0,
    }


# ---- Routes: Stripe Payments ----
@api.post("/payments/checkout")
async def create_checkout(req: CheckoutReq, request: Request, user: dict = Depends(current_user)):
    amount = int(PRICE_PACKAGES.get(req.package_id, 199) * 100)
    session = stripe.checkout.Session.create(
        payment_method_types=['card'],
        line_items=[{'price_data': {'currency': 'usd', 'product_data': {'name': 'GrantPulse Pro'}, 'unit_amount': amount}, 'quantity': 1}],
        mode='payment',
        success_url=f"{req.origin_url}/success?session_id={{CHECKOUT_SESSION_ID}}",
        cancel_url=f"{req.origin_url}/pricing",
        metadata={"user_id": user["id"]}
    )
    return {"url": session.url, "session_id": session.id}

@api.get("/payments/status/{session_id}")
async def payment_status(session_id: str, request: Request, user: dict = Depends(current_user)):
    session = stripe.checkout.Session.retrieve(session_id)
    if session.payment_status == "paid":
        await db.users.update_one({"id": user["id"]}, {"$set": {"plan": "pro"}})
    return {"status": session.status, "payment_status": session.payment_status, "amount_total": session.amount_total, "currency": session.currency}

@api.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    return {"received": True}


# ---- Scheduler: 4-hour scout ----
scheduler = AsyncIOScheduler(timezone="UTC")


async def cron_scout_all_users():
    logger.info("[cron] running 4h scout for all users")
    users = await db.users.find({"plan": "pro"}, {"_id": 0, "password": 0}).to_list(1000)
    for u in users:
        try:
            await _run_scout_for_user(u)
        except Exception as e:
            logger.error(f"[cron] user {u.get('email')} failed: {e}")


@app.on_event("startup")
async def on_startup():
    try:
        init_storage()
    except Exception as e:
        logger.warning(f"storage init deferred: {e}")
    # Indexes
    await db.users.create_index("email", unique=True)
    await db.grants.create_index([("user_id", 1), ("external_id", 1)], unique=True)
    await db.personas.create_index("user_id", unique=True)
    # Cron
    scheduler.add_job(cron_scout_all_users, "interval", hours=4, id="scout", replace_existing=True)
    scheduler.start()
    # Seed admin
    admin_email = "admin@grantpulse.io"
    if not await db.users.find_one({"email": admin_email}):
        await db.users.insert_one({
            "id": new_id(),
            "email": admin_email,
            "password": hash_password("Admin@12345"),
            "role": "admin",
            "plan": "pro",
            "free_leads_used": 0,
            "free_leads_period_start": utcnow_iso(),
            "created_at": utcnow_iso(),
            "company_name": "GrantPulse HQ",
        })
        logger.info("Seeded admin user")


@app.on_event("shutdown")
async def on_shutdown():
    try:
        scheduler.shutdown(wait=False)
    except Exception:
        pass
    client.close()


app.include_router(api)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)
